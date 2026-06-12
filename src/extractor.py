from docx import Document

def textExtractor(file_path):
    # Extracts all text from a .docx file including paragraphs and tables
    try:
        doc = Document(file_path)
        full_text = []
        
        # extract text from main paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # extract text from tables, if any(common in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        # We use strip to avoid adding empty cells
                        full_text.append(cell.text.strip())
                        
        return "\n".join(full_text)
    
    except Exception as e:
        print(f"Error: Could not read the file at {file_path}. {e}")
        return None

if __name__ == "__main__":
    # quick test logic
    resume_path = "data/input_resume.docx"
    content = textExtractor(resume_path)
    if content:
        print("Successfully extracted resume text.")
        print(content[:200] + "...") # print first 200 chars for verification
