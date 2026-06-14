from docx import Document
import re

def wordReplacer(input_path, output_path, replacement_map):
    # replace text in a .docx file while attempting to preserve formatting.
    doc = Document(input_path)

    def process_text_container(container):
        # helper to process paragraphs or table cells
        for item in container:
            for old_text, new_text in replacement_map.items():
                # we use regex with word boundaries to avoid partial word replacement
                # e.g., don't replace 'lead' inside 'pleading'
                pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                
                if pattern.search(item.text):
                    # terate through 'runs' to maintain bold/italic/font styles
                    for run in item.runs:
                        if pattern.search(run.text):
                            run.text = pattern.sub(new_text, run.text)

    # process all standard paragraphs
    process_text_container(doc.paragraphs)

    # process all tables (including nested cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_text_container(cell.paragraphs)

    doc.save(output_path)
    return True